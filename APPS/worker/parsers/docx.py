"""
DOCX parser — uses python-docx to extract text from Word documents.
"""
from typing import Optional


def parse_docx(raw_bytes: bytes) -> str:
    """
    Extract text from DOCX bytes.
    Returns plain text with paragraph breaks preserved.
    """
    try:
        from io import BytesIO
        import docx

        doc = docx.Document(BytesIO(raw_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        return "\n\n".join(paragraphs)

    except ImportError:
        print("[Parser/DOCX] python-docx not installed")
        return ""
    except Exception as exc:
        print(f"[Parser/DOCX] Failed: {exc}")
        return ""